"""Build the reliability-aware mini article in the LNCS-like example style."""

from __future__ import annotations

import tempfile
import subprocess
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = ROOT / "paper"
LOCAL_ARTICLE_DIR = ROOT.parent / "мини статья2"
TEMPLATE_DOCM = LOCAL_ARTICLE_DIR / "kopishynska_682.docm"
OUT_DOCM = LOCAL_ARTICLE_DIR / "kopishynska_small_data_crop_forecasting_reliability.docm"
FIGURES = PAPER_DIR / "figures"
REPORTS = ROOT / "reports"

CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
DOCX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"

STYLE_FALLBACKS = {
    "papertitle": "Title",
    "author": "Normal",
    "address": "Normal",
    "abstract": "Normal",
    "keywords": "Normal",
    "heading1": "Heading 1",
    "p1a": "Normal",
    "figurecaption": "Caption",
    "referenceitem": "Normal",
}


def require(path: Path) -> None:
    if not path.exists() or path.stat().st_size == 0:
        raise FileNotFoundError(path)


def style_name(doc: Document, preferred: str) -> str:
    names = {style.name for style in doc.styles}
    if preferred in names:
        return preferred
    fallback = STYLE_FALLBACKS.get(preferred, "Normal")
    return fallback if fallback in names else "Normal"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def _read_xml(zip_file: ZipFile, name: str) -> ET.Element:
    return ET.fromstring(zip_file.read(name))


def _write_xml(root: ET.Element) -> bytes:
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _content_type_part(element: ET.Element) -> str:
    return element.get("PartName", "")


def _content_type_value(element: ET.Element) -> str:
    return element.get("ContentType", "")


def _convert_docm_content_types_to_docx(root: ET.Element) -> ET.Element:
    for element in list(root):
        part_name = _content_type_part(element)
        content_type = _content_type_value(element)
        if part_name == "/word/document.xml":
            element.set("ContentType", DOCX_MAIN)
        if "vba" in part_name.lower() or "vba" in content_type.lower() or part_name.startswith("/customUI/"):
            root.remove(element)
    return root


def _remove_relationships(root: ET.Element, type_fragments: tuple[str, ...]) -> ET.Element:
    for element in list(root):
        rel_type = element.get("Type", "")
        if any(fragment in rel_type for fragment in type_fragments):
            root.remove(element)
    return root


def _docm_to_editable_docx(template_docm: Path, output_docx: Path) -> None:
    """Create a temporary docx from the macro template so python-docx can edit it."""
    with ZipFile(template_docm) as source, ZipFile(output_docx, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            name = info.filename
            if name == "[Content_Types].xml":
                root = _convert_docm_content_types_to_docx(_read_xml(source, name))
                target.writestr(info, _write_xml(root))
                continue
            if name == "_rels/.rels":
                root = _remove_relationships(_read_xml(source, name), ("ui/extensibility",))
                target.writestr(info, _write_xml(root))
                continue
            if name == "word/_rels/document.xml.rels":
                root = _remove_relationships(_read_xml(source, name), ("vbaProject",))
                target.writestr(info, _write_xml(root))
                continue
            if name.startswith("customUI/") or "vba" in name.lower() or name == "word/vbaData.xml":
                continue
            target.writestr(info, source.read(name))


def _package_generated_docx_as_docm(template_docm: Path, generated_docx: Path, output_docm: Path) -> None:
    """Save the generated document as an openable macro-enabled Word file.

    The article inherits styles and numbering from the real DOCM template. Directly
    grafting the template VBA project into a new package produces a file that
    LibreOffice cannot open, so the final package is written by LibreOffice's
    DOCM filter instead of by manual ZIP surgery.
    """
    require(template_docm)
    output_docm.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        outdir = Path(tmp)
        result = subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file://{outdir / 'lo_profile'}",
                "--headless",
                "--convert-to",
                "docm:MS Word 2007 XML VBA",
                "--outdir",
                str(outdir),
                str(generated_docx),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        converted = outdir / f"{generated_docx.stem}.docm"
        if not converted.exists():
            raise RuntimeError(f"LibreOffice did not create DOCM output. stdout={result.stdout} stderr={result.stderr}")
        converted.replace(output_docm)


def set_headers(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)
        pg_num_type.set(qn("w:start"), "1")
        for paragraph in section.first_page_header.paragraphs:
            paragraph.text = ""
        for index, paragraph in enumerate(section.header.paragraphs):
            paragraph.text = "Baseline-first crop yield forecasting" if index == 0 else ""
        for index, paragraph in enumerate(section.even_page_header.paragraphs):
            paragraph.text = "O. Kopishynska et al." if index == 0 else ""


def add_p(doc: Document, text: str = "", style: str = "p1a"):
    return doc.add_paragraph(text, style=style_name(doc, style))


def add_h(doc: Document, text: str) -> None:
    add_p(doc, text, "heading1")


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(7)


def add_fig(doc: Document, image: Path, caption: str, width: float = 4.8) -> None:
    require(image)
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    add_p(doc, caption, "figurecaption")


def add_forecast_card_table(doc: Document, cards: pd.DataFrame) -> None:
    add_p(doc, "Forecast-card summary generated by the reliability layer:", "figurecaption")
    for crop_en in ["wheat", "maize", "sunflower"]:
        card = row(cards, crop_en)
        text = (
            f"{crop_en.capitalize()}: recommended method {card['recommended_method']}; "
            f"MAE {fmt(card['recommended_mae'])} t/ha; best baseline MAE {fmt(card['best_baseline_mae'])} t/ha; "
            f"test coverage {float(card['test_coverage']):.0%}; label {card['warning_label']}."
        )
        add_p(doc, text)


def add_external_check_table(doc: Document, novelty: pd.DataFrame) -> None:
    add_p(doc, "Table 1. External regional check generated by the same baseline-first decision rule.", "figurecaption")
    display = novelty[["region", "crop", "recommended_method", "reliability_label"]].copy()
    method_short = {"FORECAST.LINEAR": "F.LINEAR", "LightGBM": "LGBM", "XGBoost": "XGB", "ElasticNet": "ENet", "ARIMA": "ARIMA"}
    label_short = {
        "baseline safer": "base",
        "within expected error": "ok",
        "outside validation error scale": "warn",
    }
    display["cell"] = display.apply(
        lambda item: f"{method_short.get(item['recommended_method'], item['recommended_method'])}/"
        f"{label_short.get(item['reliability_label'], item['reliability_label'])}",
        axis=1,
    )
    pivot = display.pivot(index="region", columns="crop", values="cell").reset_index()
    crop_order = ["wheat", "maize", "sunflower"]
    region_order = {"Poltava": 0, "Vinnytsia": 1, "Cherkasy": 2, "Ukraine": 3}
    table = doc.add_table(rows=1, cols=4)
    if "Table Grid" in {style.name for style in doc.styles}:
        table.style = "Table Grid"
    headers = ["Region", "Wheat", "Maize", "Sunflower"]
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True)
    pivot = pivot.sort_values("region", key=lambda values: values.map(region_order).fillna(99))
    for _, item in pivot.iterrows():
        cells = table.add_row().cells
        _set_cell_text(cells[0], item["region"])
        for idx, crop in enumerate(crop_order, start=1):
            _set_cell_text(cells[idx], item.get(crop, ""))
    add_p(doc, "Table labels: ok - within validation-residual band; warn - outside validation error scale; base - baseline safer.", "figurecaption")


def fmt(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def load_values() -> dict[str, pd.DataFrame]:
    paths = {
        "cards": REPORTS / "forecast_cards.csv",
        "recommended": REPORTS / "recommended_methods.csv",
        "bands": REPORTS / "prediction_band_summary.csv",
        "ablation": REPORTS / "feature_group_ablation.csv",
        "multi_cards": REPORTS / "multi_region_forecast_cards.csv",
        "region_summary": REPORTS / "region_comparison_summary.csv",
        "threshold": REPORTS / "decision_threshold_sensitivity.csv",
        "novelty": REPORTS / "novelty_evidence_table.csv",
    }
    for path in paths.values():
        require(path)
    return {name: pd.read_csv(path) for name, path in paths.items()}


def row(cards: pd.DataFrame, crop_en: str) -> pd.Series:
    return cards[cards["crop_en"] == crop_en].iloc[0]


def add_front(doc: Document, values: dict[str, pd.DataFrame]) -> None:
    cards = values["cards"]
    multi_cards = values["multi_cards"]
    wheat = row(cards, "wheat")
    maize = row(cards, "maize")
    sunflower = row(cards, "sunflower")
    ml_count = int((multi_cards["recommended_type"] == "machine_learning").sum())
    total_count = int(len(multi_cards))
    add_p(doc, "Baseline-First Crop Yield Forecasting for Farm Management Decision Support with Small Official Statistics", "papertitle")
    add_p(doc, "Olena Kopishynska1 [0000-0002-3138-7215]  Mark Fedorchenko1  Yurii Utkin1 [0000-0003-2732-4438]", "author")
    add_p(doc, "Igor Sliusar1 [0000-0003-1197-5666]  Viktor Liashenko2 [0000-0003-0177-6209]", "author")
    add_p(doc, "Alla Svitlychna3 [0000-0003-3674-5787] and Svitlana Pysarenko3 [0000-0003-4575-1417]", "author")
    add_p(doc, "1,2,3 Poltava State Agrarian University, Poltava 36003, Ukraine", "address")
    add_p(doc, "olena.kopishynska@pdau.edu.ua", "address")
    abstract = (
        "Abstract. Crop yield forecasting from small official statistics is different from forecasting with dense "
        "satellite, field, or weather datasets: the sample is short, temporal leakage is easy to introduce, and "
        "ML should not be accepted unless it beats transparent baselines. This paper presents a baseline-first "
        "and reliability-aware workflow that can serve as a forecasting module for farm management and regional "
        "advisory systems. Wheat, maize, and sunflower are evaluated in Poltava, Vinnytsia, "
        "Cherkasy, and national-level Ukraine data for 2010-2024. ElasticNet, XGBoost, and LightGBM are compared "
        "with naive lag-1, linear-trend, LINEST, and ARIMA baselines under a forward temporal design. The "
        "contribution is a decision layer that recommends ML only after it clears a practical MAE margin over "
        "the best baseline and then attaches empirical validation-residual bands, test coverage, feature-group "
        "ablation, and compact FMS-compatible forecast cards. The Poltava workflow recommends "
        f"{wheat['recommended_method']} for wheat (MAE {fmt(wheat['recommended_mae'])} t/ha), "
        f"{maize['recommended_method']} for maize (MAE {fmt(maize['recommended_mae'])} t/ha), and "
        f"{sunflower['recommended_method']} for sunflower (MAE {fmt(sunflower['recommended_mae'])} t/ha). "
        "Wheat is retained as a negative/control case because baselines remain safer in several territories. "
        f"Across the external check, ML is recommended in {ml_count} of {total_count} region-crop cases. The "
        "results show that ML can help in small official-statistics settings, but only when checked against "
        "simple baselines and reported with reliability diagnostics."
    )
    add_p(doc, abstract, "abstract")
    add_p(doc, "Keywords: crop yield forecasting, small datasets, official statistics, baseline-first selection, reliability diagnostics, machine learning", "keywords")


def add_intro(doc: Document) -> None:
    add_h(doc, "Introduction")
    for text in [
        "Digital agriculture often assumes that good forecasting requires dense data: satellite images, farm-level measurements, detailed weather grids, soil properties, and management logs. Such data are valuable, but they are not always available for regional institutions. In many real tasks the stable source is official annual agricultural statistics. These records are public and repeatable, but they are short and aggregated.",
        "The authors' previous work addressed precision-farming information systems, ERP and FMS platforms for agri-food management, data-driven monitoring, and digital skills for agronomy education [1-5]. The present study continues this decision-support line by treating forecasting as a small, auditable module that can feed farm management systems rather than as an isolated model leaderboard.",
        "This creates a practical research problem. A machine-learning model can be trained on a small annual table, but the result may be fragile. Random splitting is not appropriate because the next year is the forecasting target. Hyperparameter tuning can leak future information if it is not constrained. Even when a model gives a good score, it may not be better than a simple linear trend or a lag-1 rule.",
        "The study is therefore framed as a farm-management decision-support workflow. The central question is no longer which model has the lowest test MAE. The central question is which method should be recommended after baseline comparison and reliability checks, and how that recommendation should be communicated to a regional analyst.",
        "This change is important for the scientific story. In small-data forecasting, a negative result can be as useful as a positive one. Wheat is the main example in this project: the selected ML model is not recommended because the best simple baseline remains safer. Maize and sunflower show the opposite case, where ML provides enough improvement to be useful.",
        "A second motivation is reproducibility. Forecasting papers often describe the model and final accuracy but leave operational choices implicit: which years were used for tuning, whether baselines used the same information, and how to interpret a new-year error larger than expected. The repository makes these choices explicit and keeps the generated outputs as CSV and Markdown files.",
        "The proposed workflow is intentionally modest. It does not claim that annual official statistics are enough for production forecasting at farm level. It is better understood as a transparent benchmark layer that can be run before more expensive data collection is planned. If a simple baseline is already competitive, then extra modelling complexity has to justify itself. If an ML model wins, then its recommendation is still accompanied by a reliability warning rather than presented as a certain forecast.",
        "This makes the work relevant for regional universities and public agencies. They may not control satellite archives or field sensors, but they often have official statistics and need a reproducible way to compare forecasting options. A small, auditable workflow is useful when the sample contains only fifteen annual observations per crop.",
    ]:
        add_p(doc, text)


def add_novelty(doc: Document, values: dict[str, pd.DataFrame]) -> None:
    add_h(doc, "Novelty and Contribution")
    cards = values["multi_cards"]
    region_summary = values["region_summary"]
    threshold = values["threshold"]
    ml_count = int((cards["recommended_type"] == "machine_learning").sum())
    baseline_count = int((cards["recommended_type"] == "baseline").sum())
    switch_count = (
        threshold.groupby(["region_slug", "crop"])["recommended_type"]
        .nunique()
        .gt(1)
        .sum()
    )
    for text in [
        "The novelty of the project is not that spreadsheet functions outperform machine learning, and not that one algorithm wins on all crops. Spreadsheet-style baselines are used as control models. Their role is to make the recommendation accountable: if an ML model cannot beat a transparent method by a practical margin, it should not be presented as the recommended forecast.",
        "Compared with earlier platform-oriented and data-driven decision-support work [3, 4], the new repository adds a decision layer around the models. The output is no longer only a leaderboard of ElasticNet, XGBoost, and LightGBM. It produces recommended methods, reliability labels, empirical prediction bands, feature-group diagnostics, and FMS-ready forecast cards.",
        f"The multi-region extension strengthens the contribution. The same rule is applied to Poltava, Vinnytsia, Cherkasy, and Ukraine for wheat, maize, and sunflower. This creates {len(cards)} region-crop decisions: {ml_count} recommend ML and {baseline_count} keep a baseline. The national Ukraine series is treated as a scale check rather than as a direct oblast replicate.",
        f"The threshold-sensitivity file tests practical margins of 0.00, 0.03, 0.05, and 0.10 t/ha. In the current run, {switch_count} region-crop decisions change when the margin changes. This is useful evidence for the article because it shows that the recommendation is controlled by an explicit rule, not by a manually chosen story after looking at the results.",
        "The strongest conceptual difference is the treatment of negative cases. Wheat remains a control case in Poltava, Vinnytsia, and Ukraine because the transparent baseline is safer; Cherkasy wheat is different and clears the ML rule. This mixed pattern is more credible than saying that ML is always better.",
        "Maize provides the clearest positive signal. ML is recommended for maize in all four territories, but the reliability label still matters. Poltava maize improves MAE but falls outside the validation-residual scale too often, while Vinnytsia, Cherkasy, and Ukraine are labelled within expected error. This makes the article's claim precise: ML can earn recommendation, but the recommendation is conditional.",
    ]:
        add_p(doc, text)


def add_methods(doc: Document) -> None:
    add_h(doc, "Materials and Methods")
    for text in [
        "The main case study uses official AgroStats records for Poltava region, Ukraine, from 2010 to 2024. Vinnytsia, Cherkasy, and national-level Ukraine records are added as an external check. The target crops are wheat, maize, and sunflower. The raw data are stored as separate CSV files by crop and indicator. The pipeline loads these files, harmonises crop names and indicators, converts units, and builds one annual crop-level table per territory.",
        "Yield is converted to tonnes per hectare. Fertiliser indicators are represented as comparable per-hectare or share variables where possible. Irrigation volume is converted into cubic metres per hectare and millimetres. The main feature set uses lagged values and five-year historical means, so the model receives only information that could be known before the forecasted harvest year.",
        "The temporal design is fixed. Years up to 2018 form the initial training period, 2019-2021 are used for validation and hyperparameter selection, and 2022-2024 are held out for the main test. For each forecast year, preprocessing is fitted only on previous years. This prevents imputation, scaling, or parameter choices from using future data.",
        "The ML models are ElasticNet, XGBoost, and LightGBM. They are deliberately tuned on a constrained grid because the dataset is too small for broad automated search. High-capacity neural or deep-learning architectures are not trained on the 15-year annual series, because such models would be poorly identified without denser FMS logs, sensor streams, weather grids, or satellite observations. Baselines include naive lag-1, linear trend, LINEST with lagged predictors, and ARIMA. Forecast accuracy is reported as MAE, RMSE, and MAPE, with MAE used as the main selection metric.",
        "The code is organised as a reproducible pipeline. Running `python run_all.py` loads raw data, builds features, validates the Poltava dataset, trains models, computes baselines, runs reliability diagnostics, and exports report tables. Running `python scripts/run_multi_region.py` repeats the same decision workflow for Poltava, Vinnytsia, Cherkasy, and Ukraine. Article values are taken from generated CSV files, not manual calculations.",
        "Feature construction follows a strict time direction. For each crop-year, lagged yield and lagged input variables are used instead of contemporaneous values. Rolling means are based on earlier years only. This is a small detail technically, but it is central for the interpretation of the results. Without it, the model could learn from information that would not be known at the time of making the forecast.",
        "Baselines are implemented as first-class models, not as an afterthought. The naive lag-1 rule is a persistence forecast. The linear trend baseline uses time as the only predictor. LINEST represents a spreadsheet-friendly regression option that is easy for non-programming users to reproduce. ARIMA represents a standard univariate time-series competitor. This mix is important because a practical user may prefer a weaker but transparent method if the ML gain is small.",
    ]:
        add_p(doc, text)


def add_workflow(doc: Document, values: dict[str, pd.DataFrame]) -> None:
    add_h(doc, "Reliability-Aware Forecasting Workflow")
    cards = values["cards"]
    wheat = row(cards, "wheat")
    maize = row(cards, "maize")
    sunflower = row(cards, "sunflower")
    for text in [
        "The new decision layer starts from a conservative rule: recommend ML only if it beats the best simple baseline by a practical MAE margin. In this project the margin is 0.05 t/ha. If ML does not clear the margin, the output label is `baseline safer`. This makes the wheat case explicit instead of hiding it inside a general model ranking.",
        "The second layer is reliability labelling. For the selected ML model, the workflow calculates the 80th percentile of absolute validation residuals and places a symmetric empirical band around test predictions. This is not a formal confidence interval. It is a small-sample diagnostic that asks whether test-year errors are within the recent validation error scale.",
        "The third layer is feature stability. Feature-group ablation removes whole groups such as yield history, nitrogen, mineral-treated share, or irrigation. If removing a group increases MAE, that group is treated as useful for the selected forecast. This is easier to explain to agronomic users than a large list of individual model features.",
        "The final output is a forecast card for each crop. A card reports the recommended method, recommended MAE, best ML MAE, best baseline MAE, ML gain or loss, test coverage inside the empirical band, the most useful feature group, and a short interpretation. This is the project output that would be easiest for a regional analyst to read.",
        f"In the current run, wheat receives `{wheat['warning_label']}` and is assigned to {wheat['recommended_method']}. Maize receives `{maize['warning_label']}` and is assigned to {maize['recommended_method']}. Sunflower receives `{sunflower['warning_label']}` and is assigned to {sunflower['recommended_method']}. These labels are deliberately simple because the dataset is small.",
        "The empirical bands are based only on validation residuals. This prevents the test period from defining its own uncertainty. The band width is the recent error scale observed before the final evaluation period. If a test observation falls outside that band, the problem is not automatically model failure, but it is a warning that the new year behaves differently from the validation period.",
        "The warning labels are designed for conservative use. `Within expected error` means the selected method clears the baseline rule and the test errors mostly stay inside the validation band. `Outside validation error scale` means that ML wins by MAE but the residual pattern is less reliable. `Baseline safer` means the model has not earned a recommendation over the best simple alternative.",
        "Forecast cards deliberately combine numbers and text. The card format connects the recommendation to the baseline comparison, empirical band, and most useful feature group. This makes the output closer to a farm-management decision-support object than to a conventional model-ranking table.",
    ]:
        add_p(doc, text)
    add_forecast_card_table(doc, cards)
    add_p(doc, "The labels are diagnostic warnings for a small benchmark dataset, not production-grade guarantees.")


def add_results(doc: Document, values: dict[str, pd.DataFrame]) -> None:
    add_h(doc, "Experimental Results and Discussion")
    cards = values["cards"]
    multi_cards = values["multi_cards"]
    region_summary = values["region_summary"]
    novelty = values["novelty"]
    wheat = row(cards, "wheat")
    maize = row(cards, "maize")
    sunflower = row(cards, "sunflower")
    ml_count = int((multi_cards["recommended_type"] == "machine_learning").sum())
    baseline_count = int((multi_cards["recommended_type"] == "baseline").sum())
    add_fig(doc, FIGURES / "manuscript_figure1.png", "Fig. 1. Long-term trends of yield, crop area, fertiliser use, and irrigation in Poltava region, 2010-2024.", width=4.9)
    for text in [
        "The descriptive trends show that the three crops have different scales and dynamics. Maize has the largest yield range and the largest absolute errors. Sunflower has a narrower range, so low MAE values must still be interpreted together with percentage error. Wheat is comparatively stable but still hard to forecast because trend and year-to-year variation coexist in a short series.",
        f"The baseline-first selector changes the interpretation of the results. For wheat, the best ML model has MAE {fmt(wheat['best_ml_mae'])} t/ha, while the best baseline has MAE {fmt(wheat['best_baseline_mae'])} t/ha. The recommendation is therefore {wheat['recommended_method']}, not ML. This is the negative/control result of the study.",
        f"For maize, ML is more useful. The recommended method is {maize['recommended_method']} with MAE {fmt(maize['recommended_mae'])} t/ha. The best baseline has MAE {fmt(maize['best_baseline_mae'])} t/ha, so the ML gain is {fmt(maize['ml_gain_vs_baseline'])} t/ha. This improvement clears the practical margin.",
        f"For sunflower, the recommendation is also {sunflower['recommended_method']} with MAE {fmt(sunflower['recommended_mae'])} t/ha. The best baseline has MAE {fmt(sunflower['best_baseline_mae'])} t/ha. The absolute errors are small, but this should be read against the lower yield scale of the crop.",
    ]:
        add_p(doc, text)
    add_fig(doc, FIGURES / "mae_baselines_vs_ml.png", "Fig. 2. Baselines versus selected ML models on the 2022-2024 test window.", width=4.75)
    add_fig(doc, FIGURES / "manuscript_figure3.png", "Fig. 3. Actual and forecast yield for the selected lag-only ML models.", width=4.65)
    for text in [
        f"The reliability bands add a second view. Test coverage inside the empirical validation-residual band is {float(wheat['test_coverage']):.1%} for wheat, {float(maize['test_coverage']):.1%} for maize, and {float(sunflower['test_coverage']):.1%} for sunflower. These values are not calibrated probability levels. They simply show whether test errors are ordinary relative to the validation period.",
        f"The feature-group diagnostics also differ by crop. Wheat is most sensitive to {wheat['top_feature_group']}, maize to {maize['top_feature_group']}, and sunflower to {sunflower['top_feature_group']}. This supports the main argument: under small official datasets, there is no universal predictor group or universal model winner.",
        "The maize result is useful but still cautious. The selected model improves MAE, yet its reliability label warns that test errors are often outside the validation error scale. In practice this means the forecast should be used together with the diagnostic band and not as a single certain number.",
        "The wheat result is the strongest argument for the baseline-first design. A traditional paper could still report the trained ML model and discuss its predictors. This project instead says that the baseline is safer. That makes the conclusion narrower, but more honest and more useful.",
        "The baseline comparison also prevents a common interpretation error. If all attention is placed on the lowest ML score, wheat looks like a model-selection problem. In the reliability-aware view, it is a decision problem: should the model be recommended over a transparent baseline? For wheat, the answer is no.",
        "For maize, the selected ML method improves the absolute error enough to clear the practical margin, but the empirical coverage signal shows that the test years are not fully ordinary relative to validation residuals. This mixed message should be preserved in a small-data publication rather than smoothed away.",
        "For sunflower, the model result is more stable in absolute terms, but the smaller yield scale means that percentage errors remain relevant. The forecast card therefore keeps both the recommendation and the caution visible. In a practical workflow this would support a lightweight decision: use the ML forecast as the primary value, but keep the empirical band and recent baseline close for comparison.",
        "Several limitations follow directly from the design. The annual sample is short and the test set has only three years. The empirical bands are small-sample bands, not calibrated confidence intervals. The feature diagnostics are descriptive and model-dependent. These limitations define the correct use case for the repository.",
        "The current results should not be read as a final ranking of algorithms. They demonstrate an evaluation protocol. ElasticNet, XGBoost, and LightGBM represent different modelling families, but the contribution is the protocol around them: temporal validation, baseline-first recommendation, residual-band diagnostics, and forecast-card reporting.",
        "The approach is therefore conservative by design. It rewards improvement, but only when the improvement survives comparison with simple methods and remains understandable through the diagnostic outputs. This is a practical compromise between purely statistical benchmarking and the needs of applied agricultural decision support.",
    ]:
        add_p(doc, text)
    add_fig(doc, FIGURES / "mini_feature_group_ablation.png", "Fig. 4. Feature-group ablation for the selected lag-only models.", width=4.55)
    add_h(doc, "External Regional Check")
    for text in [
        f"The external check applies the same decision rule to four territories and three crops, producing {len(multi_cards)} comparable decisions. ML is recommended in {ml_count} cases, while {baseline_count} cases remain with a baseline. This pattern is important because it avoids the claim that ML automatically dominates small official-statistics datasets.",
        "The regional comparison supports the main article framing. Maize is the most consistent positive ML case: all four territories recommend an ML method, although Poltava maize keeps the warning label outside validation error scale. Wheat is mostly a negative/control case: Poltava, Vinnytsia, and Ukraine keep FORECAST.LINEAR, while Cherkasy wheat recommends LightGBM. Sunflower is mixed: Poltava and Vinnytsia recommend ML, whereas Cherkasy and Ukraine keep ARIMA.",
        "The national Ukraine series behaves more conservatively than the oblast cases. It recommends ML for maize, but keeps baselines for wheat and sunflower. This is not treated as a contradiction. National data aggregate different regional structures, so the result is framed as a scale robustness check rather than as a direct replication of oblast-level dynamics.",
        "The practical value of the external check is reproducibility. A reviewer can inspect whether the same rule gives plausible positive, negative, and mixed cases across multiple official datasets. The result is not a single success story; it is a controlled decision workflow that can say yes, no, or yes-with-warning depending on the crop and territory.",
    ]:
        add_p(doc, text)
    add_external_check_table(doc, novelty)
    region_lines = []
    for _, item in region_summary.sort_values("region_label").iterrows():
        region_lines.append(
            f"{item['region_label']}: {int(item['ml_recommendations'])} ML recommendations and "
            f"{int(item['baseline_recommendations'])} baseline recommendations; mean band coverage "
            f"{float(item['mean_test_coverage']):.1%}."
        )
    add_p(doc, " ".join(region_lines))


def add_conclusions(doc: Document) -> None:
    add_h(doc, "Conclusions")
    for text in [
        "The project shows that small official agricultural datasets can support a reproducible forecasting benchmark, but only if evaluation is conservative. The main contribution is the multi-region baseline-first and reliability-aware decision layer. It prevents ML from being recommended automatically and makes model uncertainty visible through simple diagnostics.",
        "The final recommendations are crop- and territory-specific. Poltava wheat remains better served by a transparent baseline. Poltava maize and sunflower benefit from LightGBM, but maize also receives a reliability warning. Across Poltava, Vinnytsia, Cherkasy, and Ukraine, maize is the most consistent positive ML case, while wheat and sunflower include clear baseline-safe cases.",
        "Future work should extend the same workflow to longer time periods, more crops, and carefully selected exogenous variables. Neural and deep-learning models should be revisited only when denser FMS records, sensor data, weather grids, satellite observations, or field-operation logs are available. The code is prepared for this because the recommendation tables, forecast cards, diagnostics, and article figures are generated from reproducible scripts.",
        "For publication purposes, the key point is the change in framing. The article does not claim that machine learning is always better for crop yield forecasting. It claims that, for small official datasets, a model should pass a baseline-first decision rule and should be reported with reliability diagnostics.",
        "In its present form the workflow is a publication benchmark and a methodological template, not a replacement for local agronomic judgement or operational early-warning systems. Its value is that it makes small-data forecasting claims harder to overstate and easier to reproduce.",
    ]:
        add_p(doc, text)


def add_refs(doc: Document) -> None:
    add_h(doc, "References")
    refs = [
        "Kopishynska O., Utkin Y., Galych O., Marenych M., Sliusar I. Main aspects of the creation of managing information system at the implementation of precision farming. 2020 IEEE 11th International Conference on Dependable Systems, Services and Technologies (DESSERT), Kyiv, Ukraine, 2020, pp. 404-410. doi: 10.1109/DESSERT50317.2020.9125072.",
        "Kopishynska O., Utkin Y., Sliusar I., Muravlov V., Makhmudov K., Chip L. Application of modern enterprise resource planning systems for agri-food supply chains as a strategy for reaching the level of Industry 4.0 for non-manufacturing organizations. Engineering Proceedings. 2023;40:15. doi: 10.3390/engproc2023040015.",
        "Kopishynska O., Utkin Y., Sliusar I., Galych O., Kovpak S., Liashenko V., Barabolia O. Comprehensive management of agroecosystem productivity on the platform of specialized farm management information systems. In: Callaos N., Gaile-Sarkane E., Lace N., Sanchez B., Savoie M. (eds.) Proceedings of the 28th World Multi-Conference on Systemics, Cybernetics and Informatics: WMSCI 2024, pp. 340-347. International Institute of Informatics and Cybernetics. doi: 10.54808/WMSCI2024.01.340.",
        "Kopishynska O., Utkin Y., Sliusar I., Kalashnyk O., Moroz S., Liashenko V., Fedorchenko M., Kovpak S. Smart agricultural systems: data-driven approaches to monitoring and decision support. In: Callaos N., Gaile-Sarkane E., Lace N., Sanchez B., Savoie M. (eds.) Proceedings of the 29th World Multi-Conference on Systemics, Cybernetics and Informatics: WMSCI 2025, pp. 505-512. International Institute of Informatics and Cybernetics. doi: 10.54808/WMSCI2025.01.505.",
        "Kopishynska O., Utkin Y., Lyashenko V., Barabolia O., Kalashnik O., Moroz S., Kartashova O. Information systems and technologies in agronomy and business: employers' requirements-oriented study in agricultural universities. Proceedings of the 25th World Multi-Conference on Systemics, Cybernetics and Informatics: WMSCI 2021, pp. 113-118. https://www.iiis.org/CDs2021/CD2021Summer/papers/SA745NT.pdf.",
        "Kamilaris A., Kartakoullis A., Prenafeta-Boldu F.X. A review on the practice of big data analysis in agriculture. Computers and Electronics in Agriculture. 2017;143:23-37.",
        "Jeong J.H. et al. Random forests for global and regional crop yield predictions. PLOS ONE. 2016;11:e0156571.",
        "Chlingaryan A., Sukkarieh S., Whelan B. Machine learning approaches for crop yield prediction and nitrogen status estimation in precision agriculture. Computers and Electronics in Agriculture. 2018;151:61-69.",
        "van Klompenburg T., Kassahun A., Catal C. Crop yield prediction using machine learning: a systematic literature review. Computers and Electronics in Agriculture. 2020;177:105709.",
        "Muruganantham P. et al. A systematic literature review on crop yield prediction with deep learning and remote sensing. Remote Sensing. 2022;14:1990.",
        "Khaki S., Wang L. Crop yield prediction using deep neural networks. Frontiers in Plant Science. 2019;10:621.",
        "Cao J. et al. Integrating multi-source data for rice yield prediction across China using machine learning and deep learning approaches. Agricultural and Forest Meteorology. 2021;297:108275.",
        "Peng B., Guan K., Pan M., Li Y. Benefits of seasonal climate prediction and satellite data for forecasting U.S. maize yield. Geophysical Research Letters. 2018;45:9662-9671.",
        "Meroni M. et al. Yield forecasting with machine learning and small data: what gains for grains? Agricultural and Forest Meteorology. 2021;308-309:108555.",
        "Sweet L. et al. Transdisciplinary coordination is essential for advancing agricultural modeling with machine learning. One Earth. 2025;8:101233.",
        "AgroStats portal. Regional agricultural statistics of Ukraine. https://agrostats.uhmi.org.ua. Accessed 01.12.2025.",
        "NASA POWER Project. Prediction Of Worldwide Energy Resources daily meteorological data. https://power.larc.nasa.gov.",
        "Tibshirani R. Regression shrinkage and selection via the Lasso. Journal of the Royal Statistical Society Series B. 1996;58:267-288.",
        "Zou H., Hastie T. Regularization and variable selection via the Elastic Net. Journal of the Royal Statistical Society Series B. 2005;67:301-320.",
    ]
    for ref in refs:
        add_p(doc, ref, "referenceitem")


def build() -> None:
    for path in [FIGURES / "manuscript_figure1.png", FIGURES / "mae_baselines_vs_ml.png", FIGURES / "manuscript_figure3.png", FIGURES / "mini_feature_group_ablation.png"]:
        require(path)
    require(TEMPLATE_DOCM)
    values = load_values()
    with tempfile.TemporaryDirectory() as tmp:
        work = Path(tmp) / "work.docx"
        generated = Path(tmp) / "generated.docx"
        clean_generated = Path(tmp) / "clean_generated.docx"
        _docm_to_editable_docx(TEMPLATE_DOCM, work)
        doc = Document(work)
        clear_body(doc)
        set_headers(doc)
        add_front(doc, values)
        add_intro(doc)
        add_novelty(doc, values)
        add_methods(doc)
        add_workflow(doc, values)
        add_results(doc, values)
        add_conclusions(doc)
        add_refs(doc)
        doc.core_properties.title = "Baseline-First Crop Yield Forecasting for Farm Management Decision Support"
        doc.core_properties.author = "Olena Kopishynska, Mark Fedorchenko, Yurii Utkin, Igor Sliusar, Viktor Liashenko, Alla Svitlychna, Svitlana Pysarenko"
        LOCAL_ARTICLE_DIR.mkdir(parents=True, exist_ok=True)
        doc.save(generated)
        Document(generated).save(clean_generated)
        _package_generated_docx_as_docm(TEMPLATE_DOCM, clean_generated, OUT_DOCM)
    print(f"Built {OUT_DOCM}")


if __name__ == "__main__":
    build()
